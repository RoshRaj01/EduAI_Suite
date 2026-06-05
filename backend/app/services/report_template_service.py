"""
Report Template Service
-----------------------
Pipeline: PDF template upload → structure extraction (PyMuPDF + regex)
         → field identification → MongoDB data fetch → Groq curation → DOCX output.

If Groq is unavailable the pipeline raises early so the caller can notify the teacher.
"""

import io
import re
import logging
from typing import Any, Dict, List, Optional, Tuple
from dataclasses import dataclass, field

import fitz  # PyMuPDF
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.models.student import Student
from app.models.course import Course
from app.models.submission import Submission
from app.models.assignment import Assignment
from app.models.exam import Exam
from app.models.user import User
from app.models.lesson import Lesson
from app.services.groq_service import GroqService, DEFAULT_MODEL

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Data classes for extracted template structure
# ---------------------------------------------------------------------------

@dataclass
class TemplateField:
    """A single key-value field extracted from the template."""
    label: str             # e.g. "Student Name", "Attendance"
    value: str             # original value (may be blank / placeholder)
    is_blank: bool         # True if the value is empty or a placeholder
    page: int = 0
    y_position: float = 0  # vertical position for ordering


@dataclass
class TemplateSection:
    """A logical section of the report template."""
    heading: str                         # e.g. "ACADEMIC PERFORMANCE"
    content_lines: List[str] = field(default_factory=list)
    fields: List[TemplateField] = field(default_factory=list)
    is_table: bool = False
    table_headers: List[str] = field(default_factory=list)
    table_rows: List[List[str]] = field(default_factory=list)
    page: int = 0
    y_position: float = 0


@dataclass
class ExtractedTemplate:
    """Full parsed template structure."""
    title: str = ""
    sections: List[TemplateSection] = field(default_factory=list)
    standalone_fields: List[TemplateField] = field(default_factory=list)
    raw_text: str = ""


# ---------------------------------------------------------------------------
# Regex patterns for common academic report fields
# ---------------------------------------------------------------------------

# Maps regex → canonical field name
FIELD_PATTERNS: List[Tuple[re.Pattern, str]] = [
    (re.compile(r"(?i)student\s*(?:\'s\s*)?name\s*[:=\-]?\s*(.*)"), "student_name"),
    (re.compile(r"(?i)name\s+of\s+(?:the\s+)?student\s*[:=\-]?\s*(.*)"), "student_name"),
    (re.compile(r"(?i)reg(?:istration)?\s*(?:no|number|#)\.?\s*[:=\-]?\s*(.*)"), "registration_number"),
    (re.compile(r"(?i)roll\s*(?:no|number|#)\.?\s*[:=\-]?\s*(.*)"), "registration_number"),
    (re.compile(r"(?i)department\s*[:=\-]?\s*(.*)"), "department"),
    (re.compile(r"(?i)class\s*(?:\/\s*section)?\s*[:=\-]?\s*(.*)"), "student_class"),
    (re.compile(r"(?i)section\s*[:=\-]?\s*(.*)"), "student_class"),
    (re.compile(r"(?i)attendance\s*(?:\(%?\))?\s*[:=\-]?\s*(.*)"), "attendance"),
    (re.compile(r"(?i)course\s*(?:name)?\s*[:=\-]?\s*(.*)"), "course_name"),
    (re.compile(r"(?i)batch\s*[:=\-]?\s*(.*)"), "batch"),
    (re.compile(r"(?i)teacher\s*(?:\'s\s*)?(?:name)?\s*[:=\-]?\s*(.*)"), "teacher_name"),
    (re.compile(r"(?i)semester\s*[:=\-]?\s*(.*)"), "semester"),
    (re.compile(r"(?i)academic\s*year\s*[:=\-]?\s*(.*)"), "academic_year"),
    (re.compile(r"(?i)date\s*[:=\-]?\s*(.*)"), "date"),
    (re.compile(r"(?i)email\s*[:=\-]?\s*(.*)"), "email"),
    (re.compile(r"(?i)avg(?:erage)?\s*(?:score|marks|grade)\s*[:=\-]?\s*(.*)"), "avg_score"),
    (re.compile(r"(?i)total\s*(?:score|marks)\s*[:=\-]?\s*(.*)"), "total_score"),
    (re.compile(r"(?i)grade\s*[:=\-]?\s*(.*)"), "grade"),
    (re.compile(r"(?i)parent\s*(?:\'s\s*)?(?:name)?\s*[:=\-]?\s*(.*)"), "parent_name"),
    (re.compile(r"(?i)contact\s*(?:no|number)?\s*[:=\-]?\s*(.*)"), "contact"),
]

# Section heading patterns for free-text sections that Groq should curate
FREE_TEXT_SECTIONS = {
    "strengths", "strength", "areas of strength",
    "weaknesses", "weakness", "areas for improvement", "areas of improvement",
    "recommendations", "recommendation", "suggestions",
    "overall assessment", "general assessment", "overall performance",
    "teacher comments", "teacher's comments", "teacher remarks",
    "conclusion", "summary", "remarks",
    "academic performance", "performance analysis", "performance summary",
    "attendance analysis", "engagement analysis",
    "co-curricular activities", "extra-curricular",
}

_BLANK_RE = re.compile(r"^[\s_.\-–—]*$")


def _is_blank_value(value: str) -> bool:
    """Check if a field value is effectively blank / placeholder."""
    v = value.strip()
    if not v:
        return True
    if _BLANK_RE.match(v):
        return True
    # Pure whitespace or repeated chars
    if len(set(v)) <= 2 and len(v) > 1:
        return True
    return False


# ---------------------------------------------------------------------------
# Step 1 — Extract template structure from PDF
# ---------------------------------------------------------------------------

class ReportTemplateService:
    """End-to-end pipeline: PDF template → filled DOCX report."""

    @staticmethod
    def extract_template(pdf_bytes: bytes) -> ExtractedTemplate:
        """
        Parse the uploaded PDF and extract its structural elements:
        sections, fields, tables, and raw text.
        """
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        template = ExtractedTemplate()
        all_lines: List[dict] = []  # {text, page, y, font_size, is_bold, x}

        for page_num in range(len(doc)):
            page = doc[page_num]
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]

            for block in blocks:
                if block["type"] != 0:  # text block
                    continue
                for line in block.get("lines", []):
                    line_text_parts = []
                    max_font_size = 0
                    is_bold = False
                    x_pos = line["bbox"][0]

                    for span in line.get("spans", []):
                        text = span.get("text", "")
                        line_text_parts.append(text)
                        font_size = span.get("size", 12)
                        if font_size > max_font_size:
                            max_font_size = font_size
                        flags = span.get("flags", 0)
                        # Bit 4 = bold in PyMuPDF flags
                        if flags & (1 << 4):
                            is_bold = True

                    full_text = "".join(line_text_parts).strip()
                    if not full_text:
                        continue

                    all_lines.append({
                        "text": full_text,
                        "page": page_num,
                        "y": line["bbox"][1],
                        "x": x_pos,
                        "font_size": max_font_size,
                        "is_bold": is_bold,
                    })

        doc.close()

        # Build raw text
        template.raw_text = "\n".join(l["text"] for l in all_lines)

        # --- Detect title (first large bold line) ---
        for line_info in all_lines[:10]:
            if line_info["font_size"] >= 14 or line_info["is_bold"]:
                template.title = line_info["text"]
                break

        # --- Group lines into sections ---
        current_section: Optional[TemplateSection] = None

        for line_info in all_lines:
            text = line_info["text"]
            is_heading = (
                line_info["is_bold"]
                or text.isupper() and len(text) > 3
                or line_info["font_size"] >= 13
                or text.endswith(":")
            )

            # Check if this line is a key:value field
            field_match = None
            for pattern, field_name in FIELD_PATTERNS:
                m = pattern.match(text)
                if m:
                    captured_value = m.group(1).strip() if m.lastindex else ""
                    field_match = TemplateField(
                        label=field_name,
                        value=captured_value,
                        is_blank=_is_blank_value(captured_value),
                        page=line_info["page"],
                        y_position=line_info["y"],
                    )
                    break

            if field_match:
                if current_section:
                    current_section.fields.append(field_match)
                else:
                    template.standalone_fields.append(field_match)
                continue

            if is_heading and text != template.title:
                # Start a new section
                heading = text.rstrip(":").strip()
                current_section = TemplateSection(
                    heading=heading,
                    page=line_info["page"],
                    y_position=line_info["y"],
                )
                template.sections.append(current_section)
            elif current_section:
                current_section.content_lines.append(text)

        # --- Detect simple table patterns within sections ---
        for section in template.sections:
            if not section.content_lines:
                continue
            # If multiple lines have consistent column separators (tabs or 3+ spaces)
            tab_lines = [l for l in section.content_lines if "\t" in l or "   " in l]
            if len(tab_lines) >= 2:
                section.is_table = True
                # First tab-separated line → headers
                delim = "\t" if "\t" in tab_lines[0] else "   "
                section.table_headers = [c.strip() for c in tab_lines[0].split(delim) if c.strip()]
                for row_line in tab_lines[1:]:
                    cols = [c.strip() for c in row_line.split(delim) if c.strip()]
                    if cols:
                        section.table_rows.append(cols)

        logger.info(
            f"Template extracted: title='{template.title}', "
            f"{len(template.sections)} sections, "
            f"{len(template.standalone_fields)} standalone fields"
        )
        return template

    # -------------------------------------------------------------------
    # Step 2 — Map fields to MongoDB data
    # -------------------------------------------------------------------

    @staticmethod
    async def fetch_data_for_fields(
        template: ExtractedTemplate,
        target_type: str,  # "student" or "course" or generic
        target_id: Optional[int] = None,
    ) -> Dict[str, Any]:
        """
        Query MongoDB for data that corresponds to the fields found in the
        template. Returns a dict keyed by canonical field name.
        """
        data: Dict[str, Any] = {}
        NOT_FOUND = "[DATA NOT FOUND — MANUAL ENTRY REQUIRED]"

        # --- Student data ---
        student = None
        if target_id:
            student = await Student.find_one(Student.int_id == target_id)

        if student:
            data["student_name"] = student.name or NOT_FOUND
            data["registration_number"] = student.registration_number or NOT_FOUND
            data["department"] = student.department or NOT_FOUND
            data["student_class"] = student.student_class or NOT_FOUND
            data["attendance"] = f"{student.attendance}%" if student.attendance is not None else NOT_FOUND
            data["avg_score"] = str(student.avg_score) if student.avg_score else NOT_FOUND

            # Email from User
            user = await User.find_one(User.email == student.email) if student.email else None
            data["email"] = student.email or NOT_FOUND

            # Course data via student's course_id
            course = await Course.find_one(Course.int_id == student.course_id) if student.course_id else None
            if course:
                data["course_name"] = course.name or NOT_FOUND
                data["batch"] = course.batch or NOT_FOUND
                data["teacher_name"] = course.teacher_name or NOT_FOUND
            else:
                data["course_name"] = NOT_FOUND
                data["batch"] = NOT_FOUND
                data["teacher_name"] = NOT_FOUND

            # Submission / Assignment scores
            user_obj = await User.find_one(User.email == student.email) if student.email else None
            user_id = user_obj.int_id if user_obj else -1

            submissions = await Submission.find(
                Submission.student_name == student.name
            ).to_list()

            assignment_details = []
            assignment_scores = []
            for s in submissions:
                assignment = await Assignment.find_one(Assignment.int_id == s.assignment_id)
                title = assignment.title if assignment else "Unknown"
                max_pts = (assignment.max_points if assignment and getattr(assignment, "max_points", None) else 100)
                if s.grade is not None:
                    pct = (s.grade / max_pts * 100) if max_pts > 0 else 0
                    assignment_details.append({"title": title, "score": s.grade, "max": max_pts, "pct": round(pct, 1)})
                    assignment_scores.append(pct)

            data["assignment_details"] = assignment_details
            data["avg_assignment_score"] = (
                f"{sum(assignment_scores) / len(assignment_scores):.1f}%"
                if assignment_scores else NOT_FOUND
            )

            # Exam scores
            exams_with_attempts = await Exam.find({"attempts.student_id": user_id}).to_list()
            exam_details = []
            exam_scores = []
            for exam in exams_with_attempts:
                max_score = sum(q.points for q in exam.questions) if exam.questions else 100
                for a in exam.attempts:
                    if a.student_id == user_id and a.score is not None:
                        pct = (a.score / max_score * 100) if max_score > 0 else 0
                        exam_details.append({
                            "title": exam.title or "Unknown Exam",
                            "score": a.score,
                            "max": max_score,
                            "pct": round(pct, 1),
                        })
                        exam_scores.append(pct)

            data["exam_details"] = exam_details
            data["avg_exam_score"] = (
                f"{sum(exam_scores) / len(exam_scores):.1f}%"
                if exam_scores else NOT_FOUND
            )

            # Lessons covered
            lessons = await Lesson.find(Lesson.course_id == student.course_id).to_list() if student.course_id else []
            data["lessons_covered"] = [l.topic for l in lessons if l.topic] if lessons else []
        else:
            # If no student found, try course-level data
            course = await Course.find_one(Course.int_id == target_id) if target_id else None
            if course:
                data["course_name"] = course.name or NOT_FOUND
                data["batch"] = course.batch or NOT_FOUND
                data["teacher_name"] = course.teacher_name or NOT_FOUND

                students = await Student.find(Student.course_id == course.int_id).to_list()
                data["total_students"] = len(students)
                data["avg_attendance"] = (
                    f"{sum(s.attendance for s in students) / len(students):.1f}%"
                    if students else NOT_FOUND
                )
            # Mark all unmapped fields as NOT_FOUND
            all_field_names = set()
            for f in template.standalone_fields:
                all_field_names.add(f.label)
            for sec in template.sections:
                for f in sec.fields:
                    all_field_names.add(f.label)
            for name in all_field_names:
                if name not in data:
                    data[name] = NOT_FOUND

        # Misc fields
        import datetime
        data.setdefault("date", datetime.datetime.now().strftime("%B %d, %Y"))
        data.setdefault("semester", NOT_FOUND)
        data.setdefault("academic_year", NOT_FOUND)
        data.setdefault("parent_name", NOT_FOUND)
        data.setdefault("contact", NOT_FOUND)
        data.setdefault("grade", NOT_FOUND)
        data.setdefault("total_score", NOT_FOUND)

        return data

    # -------------------------------------------------------------------
    # Step 3 — Groq content generation for free-text sections
    # -------------------------------------------------------------------

    @staticmethod
    def check_groq_available() -> bool:
        """Check if Groq is initialized and reachable."""
        return GroqService._available and GroqService._client is not None

    @staticmethod
    def generate_section_content(
        section_heading: str,
        data: Dict[str, Any],
        existing_content: str = "",
    ) -> str:
        """
        Use Groq to curate professional prose for a report section.
        Raises RuntimeError if Groq is unavailable (caller should have checked).
        """
        NOT_FOUND = "[DATA NOT FOUND — MANUAL ENTRY REQUIRED]"

        # Build a context summary from available data
        context_parts = []
        for key, value in data.items():
            if isinstance(value, list):
                continue
            if value and value != NOT_FOUND:
                label = key.replace("_", " ").title()
                context_parts.append(f"- {label}: {value}")

        # Assignment details
        if data.get("assignment_details"):
            context_parts.append("- Assignment Scores:")
            for a in data["assignment_details"]:
                context_parts.append(f"  • {a['title']}: {a['pct']}%")

        # Exam details
        if data.get("exam_details"):
            context_parts.append("- Exam Scores:")
            for e in data["exam_details"]:
                context_parts.append(f"  • {e['title']}: {e['pct']}%")

        context_str = "\n".join(context_parts) if context_parts else "No data available."

        prompt = f"""You are writing a section of a professional academic report.

SECTION HEADING: "{section_heading}"

AVAILABLE DATA:
{context_str}

EXISTING TEMPLATE CONTENT (if any):
{existing_content if existing_content.strip() else "None — this section was blank in the template."}

INSTRUCTIONS:
1. Write professional, detailed content for this section based on the available data.
2. If specific data is missing, write: "[DATA NOT FOUND — MANUAL ENTRY REQUIRED]" for that specific piece.
3. Keep the tone formal, encouraging, and constructive.
4. Do NOT invent data that is not provided above.
5. Write 3-6 sentences for this section.
6. Do NOT include the section heading in your response — just the content.

Write the content now:"""

        try:
            message = GroqService._client.chat.completions.create(
                messages=[{"role": "user", "content": prompt}],
                model=DEFAULT_MODEL,
                temperature=0.6,
                max_tokens=500,
            )
            result = message.choices[0].message.content.strip()
            return result if result else NOT_FOUND
        except Exception as e:
            logger.error(f"Groq content generation failed for section '{section_heading}': {e}")
            raise RuntimeError(f"Groq API error: {e}")

    # -------------------------------------------------------------------
    # Step 4 — DOCX generation
    # -------------------------------------------------------------------

    @staticmethod
    def build_docx(
        template: ExtractedTemplate,
        data: Dict[str, Any],
        generated_sections: Dict[str, str],
    ) -> bytes:
        """
        Create a DOCX file that mirrors the extracted template structure,
        filled with data and generated content.
        """
        NOT_FOUND = "[DATA NOT FOUND — MANUAL ENTRY REQUIRED]"
        doc = DocxDocument()

        # --- Document styles ---
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # --- Title ---
        title_text = template.title or "Academic Report"
        title_para = doc.add_heading(title_text, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph("")  # spacing

        # --- Standalone fields (top-level key:value pairs) ---
        if template.standalone_fields:
            for tf in template.standalone_fields:
                value = data.get(tf.label, tf.value if not tf.is_blank else NOT_FOUND)
                p = doc.add_paragraph()
                run_label = p.add_run(f"{tf.label.replace('_', ' ').title()}: ")
                run_label.bold = True
                run_label.font.size = Pt(11)
                run_value = p.add_run(str(value))
                run_value.font.size = Pt(11)
                if str(value) == NOT_FOUND:
                    run_value.font.color.rgb = RGBColor(200, 50, 50)

        # --- Sections ---
        for section in template.sections:
            doc.add_heading(section.heading, level=2)

            # Fields within the section
            for tf in section.fields:
                value = data.get(tf.label, tf.value if not tf.is_blank else NOT_FOUND)
                p = doc.add_paragraph()
                run_label = p.add_run(f"{tf.label.replace('_', ' ').title()}: ")
                run_label.bold = True
                run_value = p.add_run(str(value))
                if str(value) == NOT_FOUND:
                    run_value.font.color.rgb = RGBColor(200, 50, 50)

            # Table content
            if section.is_table and section.table_headers:
                num_cols = len(section.table_headers)
                table = doc.add_table(rows=1, cols=num_cols)
                table.style = "Light Grid Accent 1"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                # Header row
                for i, header in enumerate(section.table_headers):
                    cell = table.rows[0].cells[i]
                    cell.text = header
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

                # Data rows
                for row_data in section.table_rows:
                    row = table.add_row()
                    for i, cell_text in enumerate(row_data):
                        if i < num_cols:
                            row.cells[i].text = cell_text

                # If we have assignment/exam data and this looks like a scores table
                heading_lower = section.heading.lower()
                if any(kw in heading_lower for kw in ("assignment", "submission", "coursework")):
                    if data.get("assignment_details"):
                        for a in data["assignment_details"]:
                            row = table.add_row()
                            cells = row.cells
                            if num_cols >= 3:
                                cells[0].text = a["title"]
                                cells[1].text = str(a["score"])
                                cells[2].text = f"{a['pct']}%"
                            elif num_cols >= 2:
                                cells[0].text = a["title"]
                                cells[1].text = f"{a['pct']}%"

                if any(kw in heading_lower for kw in ("exam", "test", "assessment")):
                    if data.get("exam_details"):
                        for e in data["exam_details"]:
                            row = table.add_row()
                            cells = row.cells
                            if num_cols >= 3:
                                cells[0].text = e["title"]
                                cells[1].text = str(e["score"])
                                cells[2].text = f"{e['pct']}%"
                            elif num_cols >= 2:
                                cells[0].text = e["title"]
                                cells[1].text = f"{e['pct']}%"

            # Free-text / generated content
            elif section.heading.lower().strip() in FREE_TEXT_SECTIONS or section.heading in generated_sections:
                content = generated_sections.get(
                    section.heading,
                    "\n".join(section.content_lines) if section.content_lines else NOT_FOUND,
                )
                p = doc.add_paragraph(content)
                p.style = doc.styles["Normal"]
                if content == NOT_FOUND:
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(200, 50, 50)

            # Regular content lines (non-table, non-free-text)
            elif section.content_lines and not section.is_table:
                for line in section.content_lines:
                    doc.add_paragraph(line)

        # --- Footer ---
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run("— Report generated by EduAI Suite —")
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

        # Serialize to bytes
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()

    # -------------------------------------------------------------------
    # Full pipeline
    # -------------------------------------------------------------------

    @classmethod
    async def generate_report_from_template(
        cls,
        pdf_bytes: bytes,
        target_id: Optional[int] = None,
    ) -> bytes:
        """
        Full pipeline: PDF bytes in → DOCX bytes out.
        Raises RuntimeError if Groq is not available.
        """
        # 0. Pre-check Groq availability
        if not cls.check_groq_available():
            raise RuntimeError(
                "Groq AI service is currently unavailable. "
                "Cannot generate report content. Please try again later."
            )

        # 1. Extract template structure
        logger.info("Step 1/4: Extracting template structure from PDF...")
        template = cls.extract_template(pdf_bytes)

        # 2. Fetch data from MongoDB
        logger.info("Step 2/4: Fetching data from MongoDB...")
        data = await cls.fetch_data_for_fields(template, "auto", target_id)

        # 3. Generate content for free-text sections via Groq
        logger.info("Step 3/4: Generating content via Groq AI...")
        generated_sections: Dict[str, str] = {}
        for section in template.sections:
            heading_lower = section.heading.lower().strip()
            if heading_lower in FREE_TEXT_SECTIONS:
                try:
                    existing = "\n".join(section.content_lines)
                    content = cls.generate_section_content(
                        section.heading, data, existing
                    )
                    generated_sections[section.heading] = content
                except RuntimeError:
                    # Groq failed mid-generation — mark section
                    generated_sections[section.heading] = (
                        "[CONTENT GENERATION FAILED — MANUAL ENTRY REQUIRED]"
                    )

        # 4. Build DOCX
        logger.info("Step 4/4: Building DOCX document...")
        docx_bytes = cls.build_docx(template, data, generated_sections)

        logger.info("Report generation complete!")
        return docx_bytes
