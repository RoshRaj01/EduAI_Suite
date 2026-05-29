from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File
from typing import List, Optional
from app.models.exam import Exam, ExamQuestion, ExamChoice, ExamAttempt, ExamAnswer
from app.schemas.exam import ExamCreate, ExamResponse, ExamAttemptCreate, ExamAttemptResponse, ExamAttemptSubmit, ExamAttemptDetailResponse, ExamReviewResponse
from app.utils.auth import get_current_user
from app.models.user import User
from app.models.student import Student
import PyPDF2
import docx
import io
import re
import random
from datetime import datetime

exam_router = APIRouter(prefix="/exams", tags=["Exams"])

def format_exam_response(exam: Exam) -> dict:
    data = exam.model_dump()
    data["id"] = exam.int_id
    for q in data.get("questions", []):
        q["id"] = q.get("int_id", 0)
        for c in q.get("choices", []):
            c["id"] = c.get("int_id", 0)
    return data



@exam_router.get("/stats")
async def get_exam_stats():
    total_exams = await Exam.find_all().count()
    
    # Submissions today
    today = datetime.utcnow().replace(hour=0, minute=0, second=0, microsecond=0)
    
    from app.models.student import Student
    
    exams = await Exam.find_all().to_list()
    
    submissions_today = 0
    total_attempts = 0
    total_expected_attempts = 0
    
    for exam in exams:
        student_count = await Student.find(Student.course_id == exam.course_id).count()
        total_expected_attempts += student_count if student_count > 0 else 1
        
        for attempt in exam.attempts:
            if attempt.status == "submitted":
                total_attempts += 1
                if attempt.end_time and attempt.end_time >= today:
                    submissions_today += 1
    
    if total_expected_attempts > 0:
        avg = min(100.0, (total_attempts / total_expected_attempts) * 100)
    else:
        avg = 0.0
        
    avg_completion = f"{avg:.1f}%"
    
    return {
        "total_exams": total_exams,
        "submissions_today": submissions_today,
        "avg_completion": avg_completion,
        "pending_ai_review": 0 # Placeholder for now
    }

@exam_router.get("/", response_model=List[ExamResponse])
async def get_all_exams(course_id: Optional[int] = None, current_user: User = Depends(get_current_user)):
    if current_user.role == "student":
        # Find the student record to get the course_id
        student = await Student.find_one(Student.email == current_user.email)
        if not student:
            return []
        # Students only see published exams for their course
        exams = await Exam.find(
            Exam.course_id == student.course_id,
            Exam.status == "published"
        ).to_list()
        
        valid_exams = []
        for e in exams:
            completed_attempts = sum(1 for a in e.attempts if a.student_id == current_user.int_id and a.status != "in_progress")
            if completed_attempts < e.attempts_allowed:
                valid_exams.append(e)
                
        return [ExamResponse(**format_exam_response(e)) for e in valid_exams]
    
    # Teachers/Admins can see all or filter by course_id
    query = Exam.find_all()
    if course_id:
        query = query.find(Exam.course_id == course_id)
        
    exams = await query.to_list()
    return [ExamResponse(**format_exam_response(e)) for e in exams]

@exam_router.post("/", response_model=ExamResponse)
async def create_exam(exam_data: ExamCreate):
    new_exam = Exam(
        course_id=exam_data.course_id,
        title=exam_data.title,
        description=exam_data.description,
        time_limit=exam_data.time_limit,
        attempts_allowed=exam_data.attempts_allowed,
        randomize_questions=exam_data.randomize_questions,
        status=exam_data.status
    )
    await new_exam.assign_id()

    question_id_counter = 1
    choice_id_counter = 1
    
    for i, q_data in enumerate(exam_data.questions):
        new_question = ExamQuestion(
            int_id=question_id_counter,
            question_text=q_data.question_text,
            question_type=q_data.question_type,
            points=q_data.points,
            order=q_data.order or i
        )
        question_id_counter += 1

        for choice_data in q_data.choices:
            new_choice = ExamChoice(
                int_id=choice_id_counter,
                choice_text=choice_data.choice_text,
                is_correct=choice_data.is_correct
            )
            choice_id_counter += 1
            new_question.choices.append(new_choice)
            
        new_exam.questions.append(new_question)
    
    await new_exam.insert()
    return ExamResponse(**format_exam_response(new_exam))

@exam_router.get("/course/{course_id}", response_model=List[ExamResponse])
async def get_course_exams(course_id: int):
    exams = await Exam.find(Exam.course_id == course_id).to_list()
    return [ExamResponse(**format_exam_response(e)) for e in exams]

@exam_router.get("/{exam_id}", response_model=ExamResponse)
async def get_exam(exam_id: int):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    return ExamResponse(**format_exam_response(exam))

@exam_router.post("/{exam_id}/start", response_model=ExamAttemptResponse)
async def start_exam_attempt(exam_id: int, current_user: User = Depends(get_current_user)):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    # Check for existing in-progress attempt to allow resuming
    existing_attempt = next((a for a in exam.attempts if a.student_id == current_user.int_id and a.status == "in_progress"), None)
    
    if existing_attempt:
        res = existing_attempt.model_dump()
        res["id"] = existing_attempt.int_id
        res["exam_id"] = exam.int_id
        return res
    
    # Check completed attempts
    completed_attempts = sum(1 for a in exam.attempts if a.student_id == current_user.int_id and a.status != "in_progress")
    
    if completed_attempts >= exam.attempts_allowed:
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Maximum attempts reached")

    attempt_id = max([a.int_id for a in exam.attempts] + [0]) + 1
    
    new_attempt = ExamAttempt(
        int_id=attempt_id,
        student_id=current_user.int_id,
        status="in_progress"
    )
    
    exam.attempts.append(new_attempt)
    await exam.save()
    
    res = new_attempt.model_dump()
    res["id"] = new_attempt.int_id
    res["exam_id"] = exam.int_id
    return res

@exam_router.post("/{exam_id}/attempts/{attempt_id}/submit", response_model=ExamAttemptResponse)
async def submit_exam_attempt(exam_id: int, attempt_id: int, submission: ExamAttemptSubmit):
    # Find the exam containing this attempt
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    attempt_idx = next((i for i, a in enumerate(exam.attempts) if a.int_id == attempt_id), -1)
    if attempt_idx == -1:
        raise HTTPException(status_code=404, detail="Attempt not found")
        
    attempt = exam.attempts[attempt_idx]
    
    if attempt.status != "in_progress":
        raise HTTPException(status_code=400, detail="Attempt already submitted or invalid")

    total_score = 0.0
    answer_id_counter = max([ans.int_id for ans in attempt.answers] + [0]) + 1
    
    for ans in submission.answers:
        new_answer = ExamAnswer(
            int_id=answer_id_counter,
            attempt_int_id=attempt_id,
            question_int_id=ans.question_id,
            selected_choice_id=ans.selected_choice_id
        )
        answer_id_counter += 1
        attempt.answers.append(new_answer)
        
        # Grading
        if ans.selected_choice_id:
            question = next((q for q in exam.questions if q.int_id == ans.question_id), None)
            if question:
                choice = next((c for c in question.choices if c.int_id == ans.selected_choice_id), None)
                if choice and choice.is_correct:
                    total_score += (question.points or 0.0)

    attempt.score = total_score
    attempt.status = "submitted"
    attempt.end_time = datetime.utcnow()
    
    await exam.save()
    
    res = attempt.model_dump()
    res["id"] = attempt.int_id
    res["exam_id"] = exam.int_id
    return res

@exam_router.post("/extract")
async def extract_exam_questions(file: UploadFile = File(...)):
    contents = await file.read()
    text = ""
    
    if file.filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(io.BytesIO(contents))
        for page in reader.pages:
            text += page.extract_text() + "\n"
    elif file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(contents))
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "  ".join([cell.text.strip() for cell in row.cells])
                text_parts.append(row_text)
        text = "\n".join(text_parts)
    else:
        raise HTTPException(status_code=400, detail="Unsupported format")

    # Advanced MCQ Extraction Logic
    has_explicit_q = bool(re.search(r'(?:^|\n)\s*(?:Q|Question)\s*\d+', text, re.IGNORECASE))
    if has_explicit_q:
        q_pattern = r'(?:^|\n)\s*(?:Q(?:uestion)?\s*\d+[\.\)\:\-]?)\s*'
    else:
        q_pattern = r'(?:^|\n)\s*(?:Q(?:uestion)?\s*\d+[\.\)\:\-]?|\d+[\.\)\:\-])\s*'

    q_blocks = re.split(q_pattern, text, flags=re.IGNORECASE)
    
    questions = []
    for block in q_blocks:
        if not block.strip(): continue
        
        # Extract correct answer if present in the block
        correct_answer = ""
        ans_match = re.search(r'(?:^|\n|\s)(?:Answer|Ans)[^\w]*([A-Ea-e])', block, re.IGNORECASE)
        if ans_match:
            correct_answer = ans_match.group(1).upper()
            block = block[:ans_match.start()] + block[ans_match.end():]
        
        # Try to find where options start
        label_pattern_str = r'(?:^|\n|\s)(?:\()?([A-Ea-e])(?:(?:\)|\.)\s+|(?:\)|\.)?\s*\n)'
        opt_start_match = re.search(label_pattern_str, block, re.IGNORECASE)
        
        options = []
        if opt_start_match:
            q_text = block[:opt_start_match.start()].strip()
            options_text = block[opt_start_match.start():]
            
            label_pattern = re.compile(label_pattern_str, re.IGNORECASE)
            labels = list(label_pattern.finditer(options_text))
            
            for i, match in enumerate(labels):
                label = match.group(1).upper()
                start_idx = match.end()
                if i + 1 < len(labels):
                    end_idx = labels[i+1].start()
                else:
                    end_idx = len(options_text)
                    
                opt_text = options_text[start_idx:end_idx].strip()
                options.append({"label": label, "text": opt_text})
        else:
            q_text = block.strip()
            
        if q_text and options:
            questions.append({
                "question_text": q_text.replace('\n', '\n').strip(),
                "choices": [
                    {"choice_text": o["text"].replace('\n', ' ').strip(), "is_correct": (o["label"] == correct_answer)}
                    for o in options
                ]
            })

    return questions

@exam_router.post("/extract-answers")
async def extract_exam_answers(file: UploadFile = File(...)):
    contents = await file.read()
    text = ""
    
    if file.filename.endswith(".pdf"):
        reader = PyPDF2.PdfReader(io.BytesIO(contents))
        for page in reader.pages:
            text += page.extract_text() + "\n"
    elif file.filename.endswith(".docx"):
        doc = docx.Document(io.BytesIO(contents))
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                row_text = "  ".join([cell.text.strip() for cell in row.cells])
                text_parts.append(row_text)
        text = "\n".join(text_parts)
    else:
        raise HTTPException(status_code=400, detail="Unsupported format")

    answers = {}
    pattern = re.compile(r'(?:^|\n|\s)(?:Q|Question\s*)?(\d+)\s*[\.\-\:\)]?\s*(?:\()?([A-Ea-e])(?:\)|\.)?(?=\s|$|\n)', re.IGNORECASE)
    matches = pattern.findall(text)
    for q_num, ans in matches:
        answers[str(int(q_num))] = ans.upper()
    
    return answers

@exam_router.delete("/{exam_id}")
async def delete_exam(exam_id: int):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    await exam.delete()
    return {"message": "Exam deleted successfully"}

@exam_router.put("/{exam_id}", response_model=ExamResponse)
async def update_exam(exam_id: int, exam_data: ExamCreate):
    db_exam = await Exam.find_one(Exam.int_id == exam_id)
    if not db_exam:
        raise HTTPException(status_code=404, detail="Exam not found")
    
    # Update main fields
    db_exam.title = exam_data.title
    db_exam.description = exam_data.description
    db_exam.time_limit = exam_data.time_limit
    db_exam.attempts_allowed = exam_data.attempts_allowed
    db_exam.randomize_questions = exam_data.randomize_questions
    db_exam.status = exam_data.status
    db_exam.course_id = exam_data.course_id
    
    db_exam.questions = []

    question_id_counter = 1
    choice_id_counter = 1
    
    for i, q_data in enumerate(exam_data.questions):
        new_question = ExamQuestion(
            int_id=question_id_counter,
            question_text=q_data.question_text,
            question_type=q_data.question_type,
            points=q_data.points,
            order=q_data.order or i
        )
        question_id_counter += 1

        for choice_data in q_data.choices:
            new_choice = ExamChoice(
                int_id=choice_id_counter,
                choice_text=choice_data.choice_text,
                is_correct=choice_data.is_correct
            )
            choice_id_counter += 1
            new_question.choices.append(new_choice)
            
        db_exam.questions.append(new_question)
    
    await db_exam.save()
    return ExamResponse(**format_exam_response(db_exam))

@exam_router.get("/{exam_id}/attempts", response_model=List[ExamAttemptResponse])
async def get_exam_attempts(exam_id: int):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        return []
        
    res = []
    for a in exam.attempts:
        user = await User.find_one(User.int_id == a.student_id)
        
        attempt_res = a.model_dump()
        attempt_res["id"] = a.int_id
        attempt_res["exam_id"] = exam.int_id
        attempt_res["student_name"] = user.name if user else "Unknown"
        attempt_res["student_email"] = user.email if user else "Unknown"
        res.append(attempt_res)
        
    return res

@exam_router.get("/{exam_id}/attempts/{attempt_id}", response_model=ExamAttemptDetailResponse)
async def get_attempt_details(exam_id: int, attempt_id: int):
    exam = await Exam.find_one(Exam.int_id == exam_id)
    if not exam:
        raise HTTPException(status_code=404, detail="Exam not found")
        
    attempt = next((a for a in exam.attempts if a.int_id == attempt_id), None)
    if not attempt:
        raise HTTPException(status_code=404, detail="Attempt not found")
    
    user = await User.find_one(User.int_id == attempt.student_id)
    
    res = attempt.model_dump()
    res["id"] = attempt.int_id
    res["exam_id"] = exam.int_id
    res["student_name"] = user.name if user else "Unknown"
    res["student_email"] = user.email if user else "Unknown"
    
    formatted_exam = format_exam_response(exam)
    res["exam"] = formatted_exam
    
    formatted_answers = []
    for ans in attempt.answers:
        ans_dict = ans.model_dump()
        ans_dict["id"] = ans.int_id
        ans_dict["question_id"] = ans.question_int_id
        question = next((q for q in formatted_exam.get("questions", []) if q["id"] == ans.question_int_id), None)
        ans_dict["question"] = question
        formatted_answers.append(ans_dict)
        
    res["answers"] = formatted_answers
    
    return res
